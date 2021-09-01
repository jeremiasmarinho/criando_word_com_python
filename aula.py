from docx import Document

Document = Document()

head = Document.add_heading("O gato pepeu", 0)
paragraph_1 = Document.add_paragraph("Pepeu é o gato mimado de ana")
paragraph_2 = Document.add_paragraph("Ele gosta de ficar no mato")
paragraph_3 = Document.add_paragraph("Ana ver pepeu no mato e diz:")
paragraph_4 = Document.add_paragraph()
paragraph_4.add_run("Pepeu... Pepeu...").bold = True

paragraph_5 = Document.add_paragraph("O gato ver ana e pula no colo dela")
paragraph_6 = Document.add_paragraph("Ana dá um pulo e diz: ")
paragraph_7 = Document.add_paragraph()
paragraph_7.add_run("-Pare Pepeu! Você é muito mimado, vá para sua cama pois você não me engana.").bold = True

Document.save("aula_word.docx")